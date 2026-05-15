import io
from app.repositories import conferencia_material_repository as repo

STATUS_LABEL = {
    "CONFIRMADO":   "Confirmado",
    "SEM_MATERIAL": "Sem Material",
    "PARCIAL":      "Parcial",
}

STATUS_COR = {
    "CONFIRMADO":   "success",
    "SEM_MATERIAL": "danger",
    "PARCIAL":      "warning",
}

STATUSES_VALIDOS = list(STATUS_LABEL.keys())

MIME_PERMITIDOS = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "application/pdf": "pdf",
}


def listar_planos_por_data(data: str) -> list:
    return repo.listar_planos_por_data(data)


def status_por_data(data: str) -> dict:
    return repo.status_por_data(data)


def listar_datas() -> list:
    return repo.listar_datas_com_planos()


def arquivo_por_plano(planejamento_id: int) -> dict | None:
    return repo.arquivo_por_plano(planejamento_id)


def confirmar(planejamento_id: int, status: str, observacao: str | None,
              conferido_por: str, componentes_confirmados: list | None) -> int:
    if status not in STATUSES_VALIDOS:
        raise ValueError(f"Status inválido: {status}")
    if not conferido_por or not conferido_por.strip():
        raise ValueError("Conferido por é obrigatório")
    return repo.registrar(
        planejamento_id, status, observacao,
        conferido_por.strip(), componentes_confirmados
    )


def historico_por_plano(planejamento_id: int) -> list:
    return repo.historico_por_plano(planejamento_id)


def upload_arquivo(planejamento_id: int, file_storage, uploaded_por: str) -> dict:
    mimetype = file_storage.mimetype or ""
    filename  = file_storage.filename or ""

    if mimetype not in MIME_PERMITIDOS and not filename.lower().endswith((".docx", ".pdf")):
        raise ValueError("Formato inválido. Envie um arquivo .docx ou .pdf.")

    conteudo = file_storage.read()
    if len(conteudo) > 10 * 1024 * 1024:
        raise ValueError("Arquivo muito grande. Limite de 10 MB.")

    if mimetype == "application/pdf" or filename.lower().endswith(".pdf"):
        componentes = _parse_pdf(conteudo)
        mimetype = "application/pdf"
    else:
        componentes = _parse_docx(conteudo)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    arquivo_id = repo.salvar_arquivo(
        planejamento_id, filename, mimetype, conteudo, componentes, uploaded_por
    )
    return {"id": arquivo_id, "componentes": componentes}


def _parse_docx(conteudo: bytes) -> list:
    try:
        from docx import Document
        doc = Document(io.BytesIO(conteudo))
        items = []

        for table in doc.tables:
            for i, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                linha = " | ".join(c for c in cells if c)
                if linha and i > 0:
                    items.append(linha)

        if not items:
            for para in doc.paragraphs:
                texto = para.text.strip()
                if texto:
                    items.append(texto)

        return [{"numero": i + 1, "texto": t} for i, t in enumerate(items)]
    except Exception as e:
        raise ValueError(f"Erro ao ler o arquivo DOCX: {e}")


def _parse_pdf(conteudo: bytes) -> list:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(conteudo))
        items = []
        for page in reader.pages:
            text = page.extract_text() or ""
            for linha in text.splitlines():
                linha = linha.strip()
                if linha:
                    items.append(linha)
        return [{"numero": i + 1, "texto": t} for i, t in enumerate(items)]
    except Exception as e:
        raise ValueError(f"Erro ao ler o arquivo PDF: {e}")
