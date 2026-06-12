import csv
import io
import math
import os
from typing import Optional

from app.repositories import pci_studio_repository as repo

_COLUNA_ALIASES = {
    'designator': ['designator', 'des', 'ref', 'referencia', 'comp', 'reference', 'refdes'],
    'valor':      ['valor', 'value', 'val'],
    'package':    ['package', 'footprint', 'encapsulamento', 'pkg', 'pcb footprint'],
    'descricao':  ['descricao', 'description', 'desc', 'comment'],
    'part_number':['part_number', 'pn', 'part number', 'numero de parte', 'manufacturer pn', 'mfr pn'],
    'tipo':       ['tipo', 'type', 'mount', 'mounting'],
    'pos_x':      ['pos_x', 'x', 'mid x', 'posx', 'ref x'],
    'pos_y':      ['pos_y', 'y', 'mid y', 'posy', 'ref y'],
    'angulo':     ['angulo', 'angle', 'rotation', 'rot'],
    'lado':       ['lado', 'side', 'layer', 'board side'],
}

_TIPO_MAP = {
    'smd': 'smd', 'surface': 'smd', 'surface mount': 'smd', 'sm': 'smd',
    'pth': 'pth', 'through': 'pth', 'through hole': 'pth', 'th': 'pth',
    'conector': 'conector', 'connector': 'conector', 'conn': 'conector',
}

_TIPO_COR_PADRAO = {
    'smd': '#64748b',
    'pth': '#dc2626',
    'conector': '#374151',
    'outro': '#9ca3af',
}


def extrair_header_e_rows(f_bytes: bytes, filename: str) -> tuple[list, list]:
    ext = os.path.splitext(filename.lower())[1]
    if ext in ('.xlsx', '.xls'):
        return _extrair_excel(f_bytes)
    if ext == '.docx':
        return _extrair_docx(f_bytes)
    return _extrair_csv(f_bytes)


def detectar_colunas(header: list) -> dict:
    mapa: dict = {}
    for campo, aliases in _COLUNA_ALIASES.items():
        for col in header:
            if col.lower().strip() in aliases:
                mapa[campo] = col
                break
    return mapa


def parse_bom_arquivo(f_bytes: bytes, filename: str, mapa: dict) -> list:
    _, rows = extrair_header_e_rows(f_bytes, filename)
    return _parse_rows(rows, mapa)


def parse_bom_csv(conteudo: str, mapa: dict) -> list:
    return parse_bom_arquivo(conteudo.encode('utf-8'), 'bom.csv', mapa)


def auto_match_componentes(items: list) -> tuple[list, int]:
    packages = [i['package'] for i in items if i.get('package')]
    lib_map = repo.buscar_componentes_por_package(packages)
    matched = 0
    for item in items:
        pkg = (item.get('package') or '').upper()
        if pkg in lib_map:
            c = lib_map[pkg]
            item['componente_id'] = c['id']
            item['comp_mm'] = c.get('comp_mm')
            item['larg_mm'] = c.get('larg_mm')
            item['alt_mm'] = c.get('alt_mm')
            item['cor_hex'] = c.get('cor_hex', _TIPO_COR_PADRAO.get(item['tipo'], '#64748b'))
            matched += 1
    return items, matched


def gerar_posicoes_automaticas(items: list, board_w: float, board_h: float) -> list:
    sem_pos = [i for i in items if i.get('pos_x') is None and i.get('pos_y') is None]
    if not sem_pos:
        return items

    margem = 4.0
    espaco_x = 6.0
    espaco_y = 6.0
    cols = max(1, int((board_w - 2 * margem) / espaco_x))

    for idx, item in enumerate(sem_pos):
        col = idx % cols
        row = idx // cols
        item['pos_x'] = margem + col * espaco_x
        item['pos_y'] = margem + row * espaco_y

    return items


def _extrair_csv(f_bytes: bytes) -> tuple[list, list]:
    text = f_bytes.decode('utf-8-sig', errors='replace')
    delimiter = _sniff_csv_delimiter(text)
    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
    header = list(reader.fieldnames or [])
    rows = list(reader)
    return header, rows


def _sniff_csv_delimiter(text: str) -> str:
    try:
        return csv.Sniffer().sniff(text[:4096], delimiters=',;\t|').delimiter
    except csv.Error:
        return ','


def _extrair_excel(f_bytes: bytes) -> tuple[list, list]:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(f_bytes), read_only=True, data_only=True)
    ws = wb.active
    header = None
    rows = []
    for row in ws.iter_rows(values_only=True):
        cells = [str(c).strip() if c is not None else '' for c in row]
        if not any(cells):
            continue
        if header is None:
            header = cells
        else:
            rows.append(dict(zip(header, cells)))
    return (header or []), rows


def _extrair_docx(f_bytes: bytes) -> tuple[list, list]:
    from docx import Document
    doc = Document(io.BytesIO(f_bytes))
    if not doc.tables:
        return [], []
    table = doc.tables[0]
    header = [cell.text.strip() for cell in table.rows[0].cells]
    rows = []
    for tr in table.rows[1:]:
        cells = [cell.text.strip() for cell in tr.cells]
        if any(cells):
            rows.append(dict(zip(header, cells)))
    return header, rows


def _parse_rows(rows: list, mapa: dict) -> list:
    items = []
    for row in rows:
        designator = row.get(mapa.get('designator', ''), '').strip()
        if not designator:
            continue

        tipo_raw = row.get(mapa.get('tipo', ''), '').strip().lower()
        tipo = _TIPO_MAP.get(tipo_raw, 'smd')

        items.append({
            'designator':  designator,
            'valor':       row.get(mapa.get('valor', ''), '').strip() or None,
            'package':     row.get(mapa.get('package', ''), '').strip() or None,
            'descricao':   row.get(mapa.get('descricao', ''), '').strip() or None,
            'part_number': row.get(mapa.get('part_number', ''), '').strip() or None,
            'tipo':        tipo,
            'pos_x':       _to_float(row.get(mapa.get('pos_x', ''), '')),
            'pos_y':       _to_float(row.get(mapa.get('pos_y', ''), '')),
            'angulo':      _to_float(row.get(mapa.get('angulo', ''), '')) or 0.0,
            'lado':        _normalizar_lado(row.get(mapa.get('lado', ''), '')),
        })
    return items


def _to_float(val) -> Optional[float]:
    try:
        return float(str(val).replace(',', '.').strip())
    except (ValueError, TypeError):
        return None


def _normalizar_lado(val: str) -> str:
    v = str(val).strip().lower()
    return 'bottom' if v in ('bottom', 'bot', 'b', 'inferior', 'f.cu') else 'top'
